from docx import Document
from DataConverter import DataConverter
from enum import Enum




class Filling(Enum):
    # number of last color bits for filling
    FILL3BIT = 3
    FILL6BIT = 6
    FILL9BIT = 9



class Encoding(Enum):

    ASCII = 8
    UNICODE = 16




class DocMethod_LSB:


    def __init__(self):
        """initializes default settings of LSB method"""
        self.__converter = DataConverter()


        self.__marking_symbols = list()
        self.__marking_symbols.append(self.__converter.ListToString(self.__converter.StringToBin('-')))
        self.__marking_symbols.append(self.__converter.ListToString(self.__converter.StringToBin('+')))
        self.__marking_symbols.append(self.__converter.ListToString(self.__converter.StringToBin('$')))
        self.__marking_symbols.append(self.__converter.ListToString(self.__converter.StringToBin('&')))
        self.__marking_symbols.append(self.__converter.ListToString(self.__converter.StringToBin('#')))
        self.__marking_symbols.append(self.__converter.ListToString(self.__converter.StringToBin('*')))
        self.__marking_symbols = tuple(self.__marking_symbols)


        self.__encoding_setting = Encoding.ASCII
        self.__last_bit_setting = Filling.FILL6BIT
        self.__container_len = 0


    @property
    def Encoding(self):
        return self.__encoding_setting


    @property
    def Last_bits(self):
        return self.__last_bit_setting


    @Encoding.setter
    def Encoding(self, enc_name: str):

        if enc_name == "ASCII":
            self.__encoding_setting = Encoding.ASCII
        elif enc_name == "Unicode":
            self.__encoding_setting = Encoding.UNICODE
        else:
            print("Wrong Encoding setting")


    @Last_bits.setter
    def Last_bits(self, bits: int):

        if bits == 1:
            self.__last_bit_setting = Filling.FILL3BIT
        elif bits == 2:
            self.__last_bit_setting = Filling.FILL6BIT
        elif bits == 3:
            self.__last_bit_setting = Filling.FILL9BIT
        else:
            print("Wrong bits setting")


    @property
    def diffuse(self) -> bool:
        return self.__diffuse

    @diffuse.setter
    def diffuse(self, diffuse: bool):
        self.__diffuse = diffuse






    def Embedding (self, document: Document, bin_message: str) -> Document:
        """Embedding secret binary message to document by LSB method"""

        """______ Put marks showing that document will contain something ______"""
        if not self.HasSecretMessage(document):
            self.__MarkingDocument(document)


        """______ Taking colors of symbols ______"""
        text_colors = self.__ExtractColors(document)


        """______ Embedding with chosen number of last bits ______"""
        if self.__last_bit_setting.value == 3:

            text_colors = self.__Embed_3_bits(text_colors, bin_message)

        if self.__last_bit_setting.value == 6:

            text_colors = self.__Embed_6_bits(text_colors, bin_message)

        if self.__last_bit_setting.value == 9:

            text_colors = self.__Embed_9_bits(text_colors, bin_message)


        """______ Replace new text colors into document ______"""
        return self.__InsertColors(document, text_colors)




    def __Embed_3_bits(self, text_colors: list, bin_message: str) -> list:
        """Embed 1 bit of message to each color channel of character in document"""

        text_index = 0
        for color_index in range(1, len(text_colors)):
            if text_index < len(bin_message):
                new_color = list()

                for i in range(0, 24):
                    if (i == 7 or i == 15 or i == 23) and (text_index < len(bin_message)):
                        new_color.append(bin_message[text_index])
                        text_index = text_index + 1
                    else:
                        new_color.append(text_colors[color_index][i])
                text_colors[color_index] = self.__converter.ListToString(new_color)

            else:
                break

        return text_colors





    def __Embed_6_bits(self, text_colors: list, bin_message: str) -> list:
        """Embed 2 bits of message to each color channel of character in document"""

        text_index = 0
        for color_index in range(1, len(text_colors)):
            if text_index < len(bin_message):
                new_color = list()

                for i in range(0, 24):
                    if (i == 6 or i == 7 or i == 14 or i == 15 or i == 22 or i == 23) \
                            and (text_index < len(bin_message)):

                        new_color.append(bin_message[text_index])
                        text_index = text_index + 1
                    else:
                        new_color.append(text_colors[color_index][i])
                text_colors[color_index] = self.__converter.ListToString(new_color)

            else:
                break

        return text_colors





    def __Embed_9_bits(self, text_colors: list, bin_message: str) -> list:
        """Embed 3 bits of message to each color channel of character in document"""

        text_index = 0
        for color_index in range(1, len(text_colors)):
            if text_index < len(bin_message):
                new_color = list()

                for i in range(0, 24):
                    if (i == 5 or i == 6 or i == 7 or i == 13 or i == 14 or i == 15 or i == 21 or i == 22 or i == 23) \
                            and (text_index < len(bin_message)):

                        new_color.append(bin_message[text_index])
                        text_index = text_index + 1
                    else:
                        new_color.append(text_colors[color_index][i])
                text_colors[color_index] = self.__converter.ListToString(new_color)

            else:
                break

        return text_colors






    def __MarkingDocument (self, document: Document):
        """Adding symbol to begin to mark that document has secret message"""

        paragraph = document.paragraphs[0]
        bin_color = self.__converter.RGBtoBin(paragraph.runs[0].font.color.rgb)
        marked_symbol = ""


        """______ Selecting marking symbol according to settings  ______"""

        if self.__last_bit_setting.value == 3:

            if self.__encoding_setting.value == 8:
                marked_symbol = self.__marking_symbols[0]
            else:
                marked_symbol = self.__marking_symbols[1]

        if self.__last_bit_setting.value == 6:

            if self.__encoding_setting.value == 8:
                marked_symbol = self.__marking_symbols[2]
            else:
                marked_symbol = self.__marking_symbols[3]

        if self.__last_bit_setting.value == 9:

            if self.__encoding_setting.value == 8:
                marked_symbol = self.__marking_symbols[4]
            else:
                marked_symbol = self.__marking_symbols[5]


        """______ Embedd symbol to the color of first symbol of document  ______"""

        new_color = list()
        j = 0

        for i in range(0, 24):

            if i == 6 or i == 7 or i == 14 or i == 15 or i == 22 or i == 23:
                new_color.append(marked_symbol[j])
                j = j + 1
            else:
                new_color.append(bin_color[i])

        new_color = self.__converter.ListToString(new_color)
        paragraph.runs[0].font.color.rgb = self.__converter.BinToRGB(new_color)




    def HasSecretMessage (self, document: Document) -> bool:
        """Checking if document has marked symbol in the beginning that means that document has secret message"""

        is_marked = True

        paragraph = document.paragraphs[0]
        bin_color = self.__converter.RGBtoBin(paragraph.runs[0].font.color.rgb)
        decoded_symbol = list()

        for i in 6, 7, 14, 15, 22, 23:
            decoded_symbol.append(bin_color[i])

        decoded_symbol = self.__converter.ListToString(decoded_symbol)


        """______ Comparing first symbol to marked symbols  ______"""

        if decoded_symbol == self.__marking_symbols[0]:

            self.__last_bit_setting = Filling.FILL3BIT
            self.__encoding_setting = Encoding.ASCII

        elif decoded_symbol == self.__marking_symbols[1]:

            self.__last_bit_setting = Filling.FILL3BIT
            self.__encoding_setting = Encoding.UNICODE

        elif decoded_symbol == self.__marking_symbols[2]:

            self.__last_bit_setting = Filling.FILL6BIT
            self.__encoding_setting = Encoding.ASCII

        elif decoded_symbol == self.__marking_symbols[3]:

            self.__last_bit_setting = Filling.FILL6BIT
            self.__encoding_setting = Encoding.UNICODE

        elif decoded_symbol == self.__marking_symbols[4]:

            self.__last_bit_setting = Filling.FILL9BIT
            self.__encoding_setting = Encoding.ASCII

        elif decoded_symbol == self.__marking_symbols[5]:

            self.__last_bit_setting = Filling.FILL9BIT
            self.__encoding_setting = Encoding.UNICODE

        else:

            is_marked = False


        return is_marked




    def MaxLength(self, document: Document) -> int:
        """Calculates the max number of characters which can be hidden in document"""

        sum_of_runs = 0

        for paragraph in document.paragraphs:

            for run in paragraph.runs:
                if run.text != "": # Jump over images and empty paragraphs and none type runs
                    sum_of_runs += 1


        encoding = self.__encoding_setting.value
        bits = self.__last_bit_setting.value

        self.__container_len = int((sum_of_runs - 1) * bits / encoding)

        return self.__container_len








    def Decoding (self, document: Document) -> list:
        """Decoding secret message from document"""

        text_colors = self.__ExtractColors(document)
        bits_list = list()


        """______ Extraction last bits depending on last bit setting ______"""

        if self.__last_bit_setting.value == 3:

            for i in range(1, len(text_colors)):
                bits_list.append(text_colors[i][7])
                bits_list.append(text_colors[i][15])
                bits_list.append(text_colors[i][23])

        elif self.__last_bit_setting.value == 6:

            for i in range(1, len(text_colors)):
                bits_list.append(text_colors[i][6])
                bits_list.append(text_colors[i][7])
                bits_list.append(text_colors[i][14])
                bits_list.append(text_colors[i][15])
                bits_list.append(text_colors[i][22])
                bits_list.append(text_colors[i][23])

        else:

            for i in range(1, len(text_colors)):
                bits_list.append(text_colors[i][5])
                bits_list.append(text_colors[i][6])
                bits_list.append(text_colors[i][7])
                bits_list.append(text_colors[i][13])
                bits_list.append(text_colors[i][14])
                bits_list.append(text_colors[i][15])
                bits_list.append(text_colors[i][21])
                bits_list.append(text_colors[i][22])
                bits_list.append(text_colors[i][23])


        """______ Сomposing a binary message from the last bits ______"""

        bin_message = list()
        decoded_char = list()

        for i in range(0, self.__encoding_setting.value):
            decoded_char.append(0)
        j = 0

        for i in range(0, len(bits_list)):
            decoded_char[j] = bits_list[i]

            if j == self.__encoding_setting.value - 1:
                bin_message.append(self.__converter.ListToString(decoded_char))
                j = 0
            else:
                j += 1


        return bin_message




    def __ExtractColors(self, document: Document) -> list:
        """Extract colors from runs, convert them to bin format, build list of colors"""
        text_colors = list()
        paragraphs = document.paragraphs

        for paragraph in paragraphs:
            for run in paragraph.runs:

                if run.text != "":  # Jump over images and empty paragraphs and none type runs
                    text_colors.append(self.__converter.RGBtoBin(run.font.color.rgb))

        return text_colors




    def __InsertColors(self, document: Document, text_colors: list) -> Document:
        """Insert list of changed colors to runs of document"""

        for i in range(0, len(text_colors)):
            text_colors[i] = self.__converter.BinToRGB(text_colors[i])

        paragraphs = document.paragraphs
        color_index = 0

        for paragraph in paragraphs:
            for run in paragraph.runs:

                if run.text != "":  # Jump over images and empty paragraphs and none type runs
                    run.font.color.rgb = text_colors[color_index]
                    color_index = color_index + 1

        return document



    def Diffuse_message(self, message: str, container_len: int) -> str:
        """Spreads message to all capacity of container"""

        diffused_mes = list()
        diffused_len = container_len - (container_len % len(message))
        j = 0

        for i in range(0, diffused_len):

            diffused_mes.append(message[j])
            j += 1

            if j == len(message):

                j = 0

        return self.__converter.ListToString(diffused_mes)
















