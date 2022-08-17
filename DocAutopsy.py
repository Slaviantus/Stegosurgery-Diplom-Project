from docx import Document
from copy import deepcopy
from docx.shared import RGBColor


class DocAutopsy:


    def Autopse(self, document: Document):
        """Separates all MSword document into runs of each character"""
        paragraphs = document.paragraphs

        for i in range(0, len(paragraphs)):

            # =========  Display Paragraphs to Console =========
            print(paragraphs[i].text)
            # ==============================

            if paragraphs[i].text != "":  # Jump over images and empty paragraphs
                self.__SplitParagraph(paragraphs[i])

        self.__separated_document = document


    @property
    def SeparatedDocument(self):

        return self.__separated_document


    def __SplitParagraph(self, paragraph):
        """Splits paragraph into runs for each character"""
        parag_text = paragraph.text
        char_num = len(parag_text)

        """______ Save properties of runs to list before separation  ______"""
        runs_properties = self.__SaveRunsProperties(paragraph.runs)

        """______ Clear text in existing runs  ______"""
        for run in paragraph.runs:
            run.clear()

        """______ Put chars in existing runs  ______"""

        sum_exist_runs = len(paragraph.runs)
        for i in range(0, sum_exist_runs):
            paragraph.runs[i].text = parag_text[i]

        """______ Adding new runs and put other chars  ______"""

        for i in range(sum_exist_runs, char_num):
            paragraph.add_run(parag_text[i])

        """______ Inserting saved properties back to new runs  ______"""
        self.__InsertProperties(runs_properties, paragraph.runs)

        self.__SetDefaultColor(paragraph.runs)



    def __SaveRunsProperties(self, runs) -> list:
        """Saves properties of each run before runs separating"""
        sum_runs = len(runs)

        """______ Initialisation the list of properties  ______"""
        properties = list()
        for i in range(0, sum_runs):
            run_properties = list()
            for j in range(0, 4):
                run_properties.append(0)
            properties.append(run_properties)

        """______ Saving location of old runs into list of properties ______"""
        border_index = 0
        for i in range(0, sum_runs):
            properties[i][0] = border_index
            sum_chars = len(runs[i].text)
            border_index = border_index + sum_chars - 1
            properties[i][1] = border_index
            border_index = border_index + 1

        """______ Deep copying font and style of old runs ______"""
        for i in range(0, sum_runs):
            properties[i][2] = deepcopy(runs[i].font)       #Run font
            properties[i][3] = deepcopy(runs[i].style)      #Run style


        return properties




    def __InsertProperties(self, properties, runs):
        """Inserts saved runs properties into new separated runs"""

        for setting in properties:

            for i in range(setting[0], setting[1] + 1):
                runs[i].style = setting[3]

                runs[i].font.name = setting[2].name
                runs[i].font.italic = setting[2].italic
                runs[i].font.underline = setting[2].underline
                runs[i].font.color.rgb = setting[2].color.rgb
                runs[i].font.bold = setting[2].bold
                runs[i].font.size = setting[2].size
                runs[i].font.strike = setting[2].strike
                runs[i].font.subscript = setting[2].subscript
                runs[i].font.superscript = setting[2].superscript
                runs[i].font.double_strike = setting[2].double_strike
                runs[i].font.hidden = setting[2].hidden
                runs[i].font.emboss = setting[2].emboss
                runs[i].font.all_caps = setting[2].all_caps
                runs[i].font.complex_script = setting[2].complex_script
                runs[i].font.cs_bold = setting[2].cs_bold
                runs[i].font.cs_italic = setting[2].cs_italic
                runs[i].font.highlight_color = setting[2].highlight_color
                runs[i].font.imprint = setting[2].imprint
                runs[i].font.outline = setting[2].outline
                runs[i].font.no_proof = setting[2].no_proof
                runs[i].font.shadow = setting[2].shadow
                runs[i].font.rtl = setting[2].rtl
                runs[i].font.spec_vanish = setting[2].spec_vanish
                runs[i].font.snap_to_grid = setting[2].snap_to_grid



    def __SetDefaultColor(self, runs):
        """Putting default (black) color to run, where color is None-type"""

        default_color = RGBColor(0, 0, 0)

        for run in runs:
            if run.font.color.rgb == None:
                run.font.color.rgb = default_color


