from docx import Document
from DocAutopsy import DocAutopsy
from DocMethod_LSB import DocMethod_LSB
from DataConverter import DataConverter
from Spray import Spray



class Surgeon:

    def __init__(self):
        self.__autopsy = DocAutopsy()
        self.LSB_method = DocMethod_LSB()
        self.__converter = DataConverter()
        self.__secret_message = ""
        self.spray = Spray()
        self.__spray_active = False





    def Load_Document(self, path: str):
        """Loading document with path, which comes from GUI"""

        self.__document = Document(path)


    def Save_Document(self, path: str):
        """Saving document with path, which comes from GUI"""

        self.__document.save(path)




    def Autopse_Document(self,):
        """Autopsing (separate) document into smaller runs if its possible"""

        self.__autopsy.Autopse(self.__document)
        self.__document = self.__autopsy.SeparatedDocument




    def Prepare_message(self):
        """Converting message into bin format and diffuse message if its necessery"""

        if self.LSB_method.diffuse == True:

            length = self.Available_LSB_symbols
            self.__secret_message = self.LSB_method.Diffuse_message(self.__secret_message, length)

        self.__bin_message = self.__converter.StringToBin(self.__secret_message)

        if self.LSB_method.Encoding.value == 8:  #if ASCII is chosen
            self.__bin_message = self.__converter.AllTo8bit(self.__bin_message)

        else: #if its Unicode is chosen
            self.__bin_message = self.__converter.AllTo16bit(self.__bin_message)


        if self.__spray_active:

            self.__bin_message = self.__converter.ListToString(self.__bin_message)

            sum_of_bits = self.LSB_method.MaxLength(self.__document) * self.LSB_method.Encoding.value

            self.__bin_message = self.spray.Spray_Message(self.__bin_message, sum_of_bits)




    def Use_LSB_method(self):

        self.Prepare_message()
        self.__document = self.LSB_method.Embedding(self.__document, self.__converter.ListToString(self.__bin_message))


    def LSB_decode_document(self) -> str:

        bin_message = self.LSB_method.Decoding(self.__document)

        if self.__spray_active:

            bin_message = self.__converter.ListToString(bin_message)
            bin_message = self.spray.Decode_Sprayed_Message(bin_message)

            message = list()
            letter = list()

            for i in range(0, len(bin_message)):

                if len(letter) == self.LSB_method.Encoding.value:
                    letter = self.__converter.ListToString(letter)
                    message.append(letter)
                    letter = list()

                letter.append(bin_message[i])

            message = self.__converter.BinToString(message)
            message = self.__Clear_special_characters(message)

        else:

            message = self.__converter.BinToString(bin_message)
            message = self.__Clear_special_characters(message)


        message = self.__converter.ListToString(message)

        return message


    def __Clear_special_characters(self, message: list) -> list:
        """ Deletes special characters in noise """

        new_message = list()

        for i in range(0, len(message)):

            if message[i].isalnum() or message[i] == " ":
                new_message.append(message[i])

        return new_message





    def Has_document_something(self) -> bool:

        if self.__document.paragraphs[0].runs[0].font.color.rgb != None:

            if self.LSB_method.HasSecretMessage(self.__document):
                return True

        else:
            return  False




    @property
    def Secret_message(self):

        return self.__secret_message


    @Secret_message.setter
    def Secret_message(self, message):

        self.__secret_message = message



    @property
    def Is_spray_active(self):

        return self.__spray_active


    @Is_spray_active.setter
    def Is_spray_active(self, value: bool):

        self.__spray_active = value



    @property
    def Available_LSB_symbols(self):
        """Calculate available sum of chars for secret message for LSB method"""

        length = 0

        if self.__spray_active:

            length = int(self.spray.Sprayed_bit_size(self.LSB_method.MaxLength(self.__document) * self.LSB_method.Encoding.value) / self.LSB_method.Encoding.value)

        else:

            length = self.LSB_method.MaxLength(self.__document)

        return length




    def Log_Doc(self):

        print("================== DOC LOG ===================")


        paragraphs = self.__document.paragraphs

        for i in range(0, len(paragraphs)):

            runs = paragraphs[i].runs

            for j in range(0, len(runs)):

                run = runs[j]

                if run.text != "": # Jump over images and empty paragraphs and none type runs

                    color = self.__converter.RGBtoBin(run.font.color.rgb)

                    red = list()
                    green = list()
                    blue = list()

                    for k in range(0, len(color)):

                        if k < 8:

                            red.append(color[k])

                        elif k < 16:

                            green.append(color[k])

                        else:

                            blue.append(color[k])

                    red = self.__converter.ListToString(red)
                    green = self.__converter.ListToString(green)
                    blue = self.__converter.ListToString(blue)

                    print(run.text, red, green, blue)





